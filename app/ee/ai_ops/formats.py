# SPDX-License-Identifier: LicenseRef-FileMorph-Commercial
"""Format-aware redaction for binary document types (DOCX, XLSX).

Redaction works in layers, then proves itself against the *serialized* package —
never against the same nodes it just edited (that re-scan-what-you-touched bug
falsely certified files whose PII hid in metadata / comments / sheet names):

1. Text engine, run-aware — body / table / header / footer paragraphs (DOCX) and
   string cells (XLSX) via the deterministic text engine. A changed element may
   lose intra-element run formatting (acceptable for guaranteed redaction); clean
   elements are untouched.
2. Metadata — identity / free-text document properties (author, title, subject,
   keywords, …) are stripped unconditionally: they routinely carry PII the body
   pass never sees, and names the regex detectors can't catch.
3. Package mop-up — the saved OOXML zip is re-opened and every free-text XML part
   (comments, footnotes, shared strings, hyperlink targets, residual metadata)
   has any *contiguous* PII value replaced. Reaches parts the high-level
   python-docx / openpyxl APIs do not expose.
4. Fail-closed verification over the SERIALIZED bytes — the produced package is
   re-opened and ``detect()`` runs over every XML/.rels part's text (leaf text +
   run-text concatenation, plus raw attribute scan for hyperlink targets and
   sheet names). Any residual ⇒ the document is refused (empty bytes,
   ``verification_passed=False``). This gate does not trust the redactor's node
   coverage, so a pass means the *whole file* is clean — the conservative bias is
   deliberate: a false refusal is an annoyance, a false pass is a breach.

Documents with embedded OLE objects / sub-packages / macros are refused: their
binary parts can't be scanned, so we cannot honestly certify them clean.

PDF is intentionally absent: safe PDF redaction must delete the text layer (a
black rectangle over text is trivially removable = a breach), a separate
security-critical checkpoint. We do not ship a fake PDF redaction.
"""

from __future__ import annotations

import html
import io
import re
import zipfile
from dataclasses import dataclass

from app.ee.ai_ops.detectors import detect
from app.ee.ai_ops.redaction import redact_text, replacement_for

# Identity / free-text properties cleared on every redaction. These carry PII
# (incl. personal names the regex detectors do not catch) and are invisible to
# the body pass. Structural fields (dates, revision) are left intact.
_DOCX_CORE_PROPS = (
    "author",
    "last_modified_by",
    "title",
    "subject",
    "keywords",
    "comments",
    "category",
    "content_status",
    "identifier",
)
_XLSX_CORE_PROPS = (
    "creator",
    "lastModifiedBy",
    "title",
    "subject",
    "description",
    "keywords",
    "category",
    "contentStatus",
    "identifier",
)

# OOXML text lives in .xml / .rels parts; everything else is binary we can't scan.
_TEXT_PART = re.compile(r"\.(xml|rels)$", re.IGNORECASE)
# Leaf text between tags — excludes attributes and markup, so detection never
# sees structural numbers (RSIDs, measurements, IDs) and so cannot corrupt them.
_LEAF_TEXT_RE = re.compile(r">([^<>]+)<")
# User-text run elements (Word w:t, spreadsheet/shared-string t, DrawingML a:t).
# Concatenating these re-joins PII split across runs, and contains only user text
# — never numeric cell values — so the joined blob stays false-positive-free.
_RUN_TEXT_RE = re.compile(r"<(w:t|t|a:t)(?:\s[^>]*)?>(.*?)</\1>", re.DOTALL)
# Attribute values (double-quoted; the saved package is lxml-serialized, so all
# attributes use double quotes). PII hides here too — Word docVars (w:val=…),
# custom-property names, hyperlink targets.
_ATTR_VALUE_RE = re.compile(r'="([^"]*)"')
# Structural 8-hex tokens — rsids, paraIds, and aRGB colours (rgb="FF0000FF") —
# carry no PII, but the checksum-less PHONE detector reads all-decimal ones as
# phones, and the mop-up must not rewrite a colour into a placeholder (it would
# corrupt the file). So any attribute value that is exactly 8 hex chars is
# skipped. No structurally-validated type collides (email needs "@", IBAN ≥ 15
# chars, card 13–19 digits, IPv4 has dots); only a bare 8-digit phone placed
# *solely* in an attribute would slip, and phones are still caught in text, table
# cells, and .rels tel: targets.
_RSID_SHAPE_RE = re.compile(r"[0-9A-Fa-f]{8}")
# Sheet/defined names live here and forbid the [ ] * chars our placeholders use,
# so the mop-up skips it; openpyxl handles sheet names and the verifier guards
# anything else (fail-closed).
_XLSX_STRUCTURAL_PART = "xl/workbook.xml"
# Body parts already covered run-aware by the high-level pass — excluded from the
# detect() preview's package scan so findings are not double-counted.
_DOCX_BODY_PARTS = re.compile(r"^word/(document|header\d*|footer\d*)\.xml$", re.IGNORECASE)
_XLSX_BODY_PARTS = re.compile(r"^xl/(worksheets/.+|sharedStrings)\.xml$", re.IGNORECASE)
# Excel sheet-name constraints: <= 31 chars, none of []:*?/\.
_SHEET_NAME_FORBIDDEN = re.compile(r"[\[\]:*?/\\]")

# Decompression-bomb guard: refuse a package whose parts declare more than this
# in total uncompressed size (read from the zip directory, before any part is
# inflated by python-docx / openpyxl).
_MAX_UNCOMPRESSED = 300 * 1024 * 1024

# Embedded sub-packages / OLE objects / macros: binary, unscannable → refuse.
_EMBEDDED_MARKERS = ("embeddings/", "oleobject", "vbaproject.bin")

# Friendly location labels for findings surfaced from non-body package parts.
_PART_LABELS = (
    ("docprops/", "Metadaten"),
    ("word/comments", "Kommentar"),
    ("word/footnotes", "Fußnote"),
    ("word/endnotes", "Endnote"),
    ("xl/workbook.xml", "Tabellenblatt-Name"),
    ("/comments", "Kommentar"),
)


class DocumentReadError(Exception):
    """An uploaded document could not be parsed or safely redacted."""


@dataclass(frozen=True, slots=True)
class FormatRedactionResult:
    """Outcome of redacting a binary document. ``data`` is empty on failure."""

    data: bytes
    entities_redacted: int
    verification_passed: bool
    residual_count: int


def _finding(span, location: str) -> dict:
    return {
        "entity_type": span.entity_type,
        "value": span.value,
        "location": location,
        "confidence": span.confidence,
    }


def _part_label(name: str) -> str:
    low = name.lower()
    for needle, label in _PART_LABELS:
        if needle in low:
            return label
    return "Dokument"


def _detectable_texts(name: str, raw: str) -> list[str]:
    """Text strings from one XML part to scan for PII.

    For .rels and the XLSX workbook part, in-scope PII lives in attribute values
    (hyperlink targets, sheet/defined names) whose surrounding numbers are short
    IDs, so the whole part is safe to scan raw. Every other part is scanned by:
    leaf text segments (contiguous PII; no cross-element joins or numeric-cell
    noise), a run-text concatenation (PII split across runs), and attribute values
    (Word docVars, custom-property names, …) — skipping any value that is exactly
    8 hex chars (rsid / paraId / aRGB-colour shape) so the checksum-less phone
    detector doesn't false-fire and the mop-up doesn't corrupt a colour.
    """
    low = name.lower()
    if low.endswith(".rels") or low == _XLSX_STRUCTURAL_PART:
        return [raw]
    texts = [html.unescape(s) for s in _LEAF_TEXT_RE.findall(raw)]
    texts.extend(
        html.unescape(v) for v in _ATTR_VALUE_RE.findall(raw) if not _RSID_SHAPE_RE.fullmatch(v)
    )
    runs = "".join(html.unescape(t) for _tag, t in _RUN_TEXT_RE.findall(raw))
    if runs:
        texts.append(runs)
    return texts


# ──────────────────────────────────────────────────────────────────────────
# Package-level guards, mop-up and verification (format-agnostic)
# ──────────────────────────────────────────────────────────────────────────


def _guard_package(data: bytes) -> None:
    """Reject decompression bombs and unscannable embedded objects up front.

    Reads only the zip central directory (no inflation), so it is cheap and runs
    before python-docx / openpyxl decompress anything.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            infos = zf.infolist()
    except Exception as exc:
        raise DocumentReadError("not a valid OOXML package") from exc
    if sum(i.file_size for i in infos) > _MAX_UNCOMPRESSED:
        raise DocumentReadError("document too large when decompressed")
    for i in infos:
        low = i.filename.lower()
        if any(marker in low for marker in _EMBEDDED_MARKERS):
            raise DocumentReadError("document contains embedded objects that cannot be redacted")


def _extra_findings(data: bytes, entity_types, body_re: re.Pattern) -> list[dict]:
    """Detect PII in non-body package parts (metadata, comments, sheet names).

    Makes the /detect preview honest about PII the high-level node walk can't see.
    Body parts (already covered run-aware by the high-level pass) are excluded so
    findings are not double-counted.
    """
    findings: list[dict] = []
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for name in zf.namelist():
            if not _TEXT_PART.search(name) or body_re.match(name):
                continue
            label = _part_label(name)
            seen: set[tuple[str, str]] = set()
            for text in _detectable_texts(name, zf.read(name).decode("utf-8", "replace")):
                for span in detect(text, entity_types):
                    key = (span.entity_type, span.value)
                    if key not in seen:
                        seen.add(key)
                        findings.append(_finding(span, label))
    return findings


def _mop_up_package(data: bytes, entity_types, mode: str) -> tuple[bytes, int]:
    """Redact contiguous PII in every free-text XML part of a saved package.

    Reaches parts the high-level API does not expose (comments, footnotes, shared
    strings, hyperlink targets, residual metadata). Run-split PII it cannot match
    is caught by ``_package_residual`` → fail-closed. ``xl/workbook.xml`` is
    skipped because sheet/defined names forbid the placeholder characters.
    """
    extra = 0
    out = io.BytesIO()
    with (
        zipfile.ZipFile(io.BytesIO(data)) as zin,
        zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout,
    ):
        for info in zin.infolist():
            raw = zin.read(info.filename)
            name = info.filename
            if _TEXT_PART.search(name) and name != _XLSX_STRUCTURAL_PART:
                text = raw.decode("utf-8", "replace")
                values: set[tuple[str, str]] = set()
                for chunk in _detectable_texts(name, text):
                    for s in detect(chunk, entity_types):
                        values.add((s.value, s.entity_type))
                changed = False
                # Longest first: a short value can be a substring of a longer one;
                # replacing the long match first avoids corrupting it.
                for value, etype in sorted(values, key=lambda v: len(v[0]), reverse=True):
                    n = text.count(value)
                    if n:
                        text = text.replace(value, replacement_for(etype, value, mode))
                        extra += n
                        changed = True
                if changed:
                    raw = text.encode("utf-8")
            zout.writestr(info, raw)
    return out.getvalue(), extra


def _package_residual(data: bytes, entity_types) -> int:
    """Count PII still detectable anywhere in the serialized package.

    Scans every XML/.rels part via ``_detectable_texts`` (leaf text + run-text,
    plus raw attribute scan for .rels / workbook names). Independent of the
    high-level redactor's node coverage — this is the honest fail-closed gate:
    a zero means the *whole file* is clean, not just the nodes we knew to touch.
    """
    count = 0
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for name in zf.namelist():
            if not _TEXT_PART.search(name):
                continue
            for text in _detectable_texts(name, zf.read(name).decode("utf-8", "replace")):
                count += len(detect(text, entity_types))
    return count


# ──────────────────────────────────────────────────────────────────────────
# DOCX
# ──────────────────────────────────────────────────────────────────────────


def _docx_paragraphs(doc):
    """Yield (location_label, paragraph) for every text-bearing paragraph:
    body, table cells, and each section's header/footer."""
    for p in doc.paragraphs:
        yield ("Absatz", p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield ("Tabelle", p)
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield ("Kopfzeile", p)
        for p in section.footer.paragraphs:
            yield ("Fußzeile", p)


def _open_docx(data: bytes):
    from docx import Document

    _guard_package(data)
    try:
        return Document(io.BytesIO(data))
    except Exception as exc:  # python-docx raises PackageNotFoundError / BadZipFile / ... on junk
        raise DocumentReadError("not a valid .docx") from exc


def _strip_docx_properties(doc) -> None:
    cp = doc.core_properties
    for attr in _DOCX_CORE_PROPS:
        try:
            setattr(cp, attr, "")
        except (ValueError, TypeError):
            pass


def detect_docx(data: bytes, entity_types: tuple[str, ...] | None) -> list[dict]:
    doc = _open_docx(data)
    findings: list[dict] = []
    for label, p in _docx_paragraphs(doc):
        for span in detect(p.text, entity_types):
            findings.append(_finding(span, label))
    findings.extend(_extra_findings(data, entity_types, _DOCX_BODY_PARTS))
    return findings


def redact_docx(
    data: bytes, entity_types: tuple[str, ...] | None, mode: str
) -> FormatRedactionResult:
    doc = _open_docx(data)
    total = 0
    for _label, p in _docx_paragraphs(doc):
        if not p.text:
            continue
        res = redact_text(p.text, entity_types, mode)
        if res.entities_redacted == 0:
            continue
        if not res.verification_passed:
            return FormatRedactionResult(b"", total, False, len(res.residual))
        # Rewrite the paragraph: put the redacted text in the first run and clear
        # the rest, so the concatenated paragraph text is the redacted version.
        if p.runs:
            p.runs[0].text = res.text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(res.text)
        total += res.entities_redacted

    _strip_docx_properties(doc)

    buf = io.BytesIO()
    doc.save(buf)
    saved, extra = _mop_up_package(buf.getvalue(), entity_types, mode)
    total += extra

    # Authoritative fail-closed gate over the serialized package.
    residual = _package_residual(saved, entity_types)
    if residual:
        return FormatRedactionResult(b"", total, False, residual)
    return FormatRedactionResult(saved, total, True, 0)


# ──────────────────────────────────────────────────────────────────────────
# XLSX
# ──────────────────────────────────────────────────────────────────────────


def _open_xlsx(data: bytes):
    import openpyxl

    _guard_package(data)
    try:
        return openpyxl.load_workbook(io.BytesIO(data))
    except Exception as exc:  # openpyxl raises BadZipFile / KeyError / ... on junk
        raise DocumentReadError("not a valid .xlsx") from exc


def _strip_xlsx_properties(wb) -> None:
    props = wb.properties
    for attr in _XLSX_CORE_PROPS:
        try:
            setattr(props, attr, None)
        except (ValueError, TypeError):
            pass
    # Custom document properties (docProps/custom.xml) — name AND value can carry
    # PII; delete each so neither survives at the source (the mop-up + gate still
    # back this up). openpyxl's CustomPropertyList supports item deletion by name.
    try:
        custom = wb.custom_doc_props
        for name in [p.name for p in list(custom)]:
            del custom[name]
    except (AttributeError, KeyError, TypeError):
        pass


def _xlsx_string_cells(wb):
    """Yield (sheet_title, cell) for string cells that are not formulas."""
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, str) and value and not value.startswith("="):
                    yield (ws.title, cell)


def _redact_sheet_names(wb, entity_types, mode: str) -> int:
    """Redact PII in worksheet names, honouring Excel's constraints.

    Sheet names allow neither ``[]:*?/\\`` nor >31 chars, so the placeholder is
    sanitized to a safe form and de-duplicated against the other sheet names.
    """
    used = {ws.title for ws in wb.worksheets if not detect(ws.title, entity_types)}
    count = 0
    for ws in wb.worksheets:
        if not detect(ws.title, entity_types):
            continue
        res = redact_text(ws.title, entity_types, mode)
        safe = (_SHEET_NAME_FORBIDDEN.sub("_", res.text).strip() or "Sheet")[:31]
        n = 1
        while safe in used:
            n += 1
            suffix = f"_{n}"
            safe = f"{safe[: 31 - len(suffix)]}{suffix}"  # keep <= 31 even for n >= 100
        ws.title = safe
        used.add(safe)
        count += res.entities_redacted
    return count


def detect_xlsx(data: bytes, entity_types: tuple[str, ...] | None) -> list[dict]:
    wb = _open_xlsx(data)
    findings: list[dict] = []
    for title, cell in _xlsx_string_cells(wb):
        for span in detect(cell.value, entity_types):
            findings.append(_finding(span, f"{title}!{cell.coordinate}"))
    findings.extend(_extra_findings(data, entity_types, _XLSX_BODY_PARTS))
    return findings


def redact_xlsx(
    data: bytes, entity_types: tuple[str, ...] | None, mode: str
) -> FormatRedactionResult:
    wb = _open_xlsx(data)
    total = 0
    for _title, cell in _xlsx_string_cells(wb):
        res = redact_text(cell.value, entity_types, mode)
        if res.entities_redacted == 0:
            continue
        if not res.verification_passed:
            return FormatRedactionResult(b"", total, False, len(res.residual))
        cell.value = res.text
        total += res.entities_redacted

    total += _redact_sheet_names(wb, entity_types, mode)
    _strip_xlsx_properties(wb)

    buf = io.BytesIO()
    wb.save(buf)
    saved, extra = _mop_up_package(buf.getvalue(), entity_types, mode)
    total += extra

    # Authoritative fail-closed gate over the serialized package.
    residual = _package_residual(saved, entity_types)
    if residual:
        return FormatRedactionResult(b"", total, False, residual)
    return FormatRedactionResult(saved, total, True, 0)
