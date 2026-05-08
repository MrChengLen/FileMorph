from pathlib import Path

import pytest


def _weasyprint_works() -> bool:
    """WeasyPrint depends on native GTK/Pango libs; on Windows dev hosts
    these typically aren't installed. Linux CI and the Hetzner Dockerfile
    do install them, so the DOCX → PDF tests run there. We probe by
    actually trying to render — `import weasyprint` succeeds even when
    the native libs are missing; the failure surfaces on first use."""
    try:
        import weasyprint

        weasyprint.HTML(string="<p>probe</p>").write_pdf()
        return True
    except Exception:
        return False


_WEASYPRINT_OK = _weasyprint_works()
_skip_no_weasyprint = pytest.mark.skipif(
    not _WEASYPRINT_OK,
    reason="WeasyPrint native deps (libgobject/pango) unavailable on this host",
)


def test_txt_to_pdf(client, auth_headers, sample_txt):
    with sample_txt.open("rb") as f:
        res = client.post(
            "/api/v1/convert",
            headers=auth_headers,
            files={"file": ("sample.txt", f, "text/plain")},
            data={"target_format": "pdf"},
        )
    assert res.status_code == 200
    # PDF files start with %PDF
    assert res.content[:4] == b"%PDF"


# ── DOCX → PDF (mammoth + WeasyPrint) ────────────────────────────────────────
# The previous DOCX → PDF converter imported `docx2pdf`, which was not in
# requirements.txt and crashed at runtime on every container deployment.
# These tests pin the new mammoth + WeasyPrint pipeline and guard against a
# regression to that broken state, plus verify SSRF protection.

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _make_docx(path: Path, paragraphs: list[str]) -> Path:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path


def _make_docx_with_table(path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Report header.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Region"
    table.cell(0, 1).text = "Sales"
    table.cell(1, 0).text = "EU"
    table.cell(1, 1).text = "12345"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    return _make_docx(tmp_path / "sample.docx", ["Hello FileMorph!", "Second paragraph."])


@pytest.fixture
def sample_docx_with_table(tmp_path) -> Path:
    return _make_docx_with_table(tmp_path / "with_table.docx")


@_skip_no_weasyprint
def test_docx_to_pdf_happy_path(client, auth_headers, sample_docx):
    with sample_docx.open("rb") as f:
        res = client.post(
            "/api/v1/convert",
            headers=auth_headers,
            files={"file": ("sample.docx", f, _DOCX_MIME)},
            data={"target_format": "pdf"},
        )
    assert res.status_code == 200, res.text
    assert res.content[:5] == b"%PDF-", "output is not a PDF"
    assert len(res.content) > 1024, "PDF unexpectedly small"


@_skip_no_weasyprint
def test_docx_to_pdf_with_table(client, auth_headers, sample_docx_with_table):
    with sample_docx_with_table.open("rb") as f:
        res = client.post(
            "/api/v1/convert",
            headers=auth_headers,
            files={"file": ("with_table.docx", f, _DOCX_MIME)},
            data={"target_format": "pdf"},
        )
    assert res.status_code == 200, res.text
    assert res.content[:5] == b"%PDF-"

    from io import BytesIO

    from pypdf import PdfReader

    reader = PdfReader(BytesIO(res.content))
    extracted = "\n".join((p.extract_text() or "") for p in reader.pages)
    assert "Region" in extracted, "table header missing in PDF"
    assert "EU" in extracted, "table cell missing in PDF"


@_skip_no_weasyprint
def test_docx_to_pdf_ssrf_blocked(client, auth_headers, sample_docx, monkeypatch):
    """The DOCX → PDF pipeline must not make outbound network calls.

    mammoth inlines images as data: URIs and WeasyPrint runs with
    `_deny_url_fetcher`. Even a benign-looking DOCX must convert without
    touching the network — guards against a future refactor that drops
    `url_fetcher=`.
    """
    import socket

    def _block(self, addr, *args, **kwargs):
        raise AssertionError(f"unexpected outbound network call to {addr!r}")

    monkeypatch.setattr(socket.socket, "connect", _block)

    with sample_docx.open("rb") as f:
        res = client.post(
            "/api/v1/convert",
            headers=auth_headers,
            files={"file": ("sample.docx", f, _DOCX_MIME)},
            data={"target_format": "pdf"},
        )
    assert res.status_code == 200, res.text
    assert res.content[:5] == b"%PDF-"


def test_docx_to_txt_unchanged(client, auth_headers, sample_docx):
    """DOCX → TXT uses a separate converter (python-docx) and must keep working."""
    with sample_docx.open("rb") as f:
        res = client.post(
            "/api/v1/convert",
            headers=auth_headers,
            files={"file": ("sample.docx", f, _DOCX_MIME)},
            data={"target_format": "txt"},
        )
    assert res.status_code == 200, res.text
    text = res.content.decode("utf-8")
    assert "Hello FileMorph!" in text
    assert "Second paragraph." in text
