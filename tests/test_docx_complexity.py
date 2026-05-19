# SPDX-License-Identifier: AGPL-3.0-or-later
"""DOCX complexity detector + engine-resolution unit tests.

These tests pin the routing decision that selects between the
high-fidelity LibreOffice path and the mammoth + WeasyPrint fallback in
``app/converters/document.py``. The detector itself never runs
LibreOffice — it just probes the DOCX OPC ZIP for the features mammoth
silently drops — so the tests are pure-Python and run on every host
(including Windows dev boxes that lack GTK/Pango).
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest

from app.converters.document import (
    _detect_docx_complexity,
    _docx_is_complex,
    _resolve_office_engine,
)


# ── Fixture builders ─────────────────────────────────────────────────────────


def _make_simple_docx(path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("A simple paragraph.")
    doc.add_paragraph("A second simple paragraph.")
    doc.save(str(path))
    return path


def _inject_part(src: Path, dst: Path, part_name: str, body: bytes) -> Path:
    """Add a raw zip entry to a DOCX. Used to simulate the *presence* of
    features (footnotes.xml, header1.xml, …) without faking valid OOXML —
    the detector decides on presence, not validity."""
    with zipfile.ZipFile(src) as zin, zipfile.ZipFile(dst, "w") as zout:
        for item in zin.namelist():
            zout.writestr(item, zin.read(item))
        zout.writestr(part_name, body)
    return dst


def _rewrite_document_xml(src: Path, dst: Path, transform) -> Path:
    """Apply ``transform(document_xml_bytes) -> bytes`` to word/document.xml."""
    with zipfile.ZipFile(src) as zin, zipfile.ZipFile(dst, "w") as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "word/document.xml":
                data = transform(data)
            zout.writestr(item, data)
    return dst


# ── _detect_docx_complexity ──────────────────────────────────────────────────


def test_simple_docx_has_no_complex_features(tmp_path):
    docx = _make_simple_docx(tmp_path / "simple.docx")
    flags = _detect_docx_complexity(docx)
    assert flags == {k: False for k in flags}
    assert not _docx_is_complex(flags)


def test_detects_footnotes(tmp_path):
    simple = _make_simple_docx(tmp_path / "base.docx")
    docx = _inject_part(simple, tmp_path / "with_footnotes.docx", "word/footnotes.xml", b"<x/>")
    flags = _detect_docx_complexity(docx)
    assert flags["footnotes"]
    assert _docx_is_complex(flags)


def test_detects_endnotes(tmp_path):
    simple = _make_simple_docx(tmp_path / "base.docx")
    docx = _inject_part(simple, tmp_path / "with_endnotes.docx", "word/endnotes.xml", b"<x/>")
    flags = _detect_docx_complexity(docx)
    assert flags["endnotes"]


def test_detects_headers(tmp_path):
    simple = _make_simple_docx(tmp_path / "base.docx")
    docx = _inject_part(simple, tmp_path / "with_header.docx", "word/header1.xml", b"<x/>")
    assert _detect_docx_complexity(docx)["headers"]


def test_detects_footers(tmp_path):
    simple = _make_simple_docx(tmp_path / "base.docx")
    docx = _inject_part(simple, tmp_path / "with_footer.docx", "word/footer1.xml", b"<x/>")
    assert _detect_docx_complexity(docx)["footers"]


def test_detects_ole_embeddings(tmp_path):
    simple = _make_simple_docx(tmp_path / "base.docx")
    docx = _inject_part(
        simple, tmp_path / "with_ole.docx", "word/embeddings/oleObject1.bin", b"\x00\x01\x02"
    )
    assert _detect_docx_complexity(docx)["ole"]


def test_detects_multi_section(tmp_path):
    simple = _make_simple_docx(tmp_path / "base.docx")
    # Two <w:sectPr> elements → multi-section layout.
    docx = _rewrite_document_xml(
        simple,
        tmp_path / "multi_section.docx",
        lambda xml: xml.replace(b"<w:sectPr", b"<w:sectPr/><w:sectPr", 1),
    )
    assert _detect_docx_complexity(docx)["sections"]


def test_detects_equations(tmp_path):
    simple = _make_simple_docx(tmp_path / "base.docx")
    docx = _rewrite_document_xml(
        simple,
        tmp_path / "equations.docx",
        lambda xml: xml.replace(b"<w:body>", b"<w:body><m:oMath><m:r/></m:oMath>", 1),
    )
    assert _detect_docx_complexity(docx)["equations"]


def test_detects_multilevel_lists(tmp_path):
    simple = _make_simple_docx(tmp_path / "base.docx")
    docx = _rewrite_document_xml(
        simple,
        tmp_path / "multilevel.docx",
        lambda xml: xml.replace(
            b"<w:body>",
            b'<w:body><w:p><w:pPr><w:numPr><w:ilvl w:val="2"/></w:numPr></w:pPr></w:p>',
            1,
        ),
    )
    assert _detect_docx_complexity(docx)["multilevel_lists"]


def test_corrupt_docx_returns_no_complexity(tmp_path):
    bad = tmp_path / "not_a_zip.docx"
    bad.write_bytes(b"this is not a zipfile")
    flags = _detect_docx_complexity(bad)
    assert flags == {k: False for k in flags}
    assert not _docx_is_complex(flags)


# ── _resolve_office_engine ───────────────────────────────────────────────────


def test_engine_auto_simple_doc_picks_mammoth():
    engine, warnings = _resolve_office_engine("auto", docx_is_complex=False, soffice_available=True)
    assert engine == "mammoth"
    assert warnings == []


def test_engine_auto_complex_doc_with_lo_picks_libreoffice():
    engine, warnings = _resolve_office_engine("auto", docx_is_complex=True, soffice_available=True)
    assert engine == "libreoffice"
    assert warnings == []


def test_engine_auto_complex_doc_without_lo_falls_back_with_warning():
    engine, warnings = _resolve_office_engine("auto", docx_is_complex=True, soffice_available=False)
    assert engine == "mammoth"
    assert "engine=mammoth_fallback" in warnings
    assert "reason=soffice_unavailable" in warnings


def test_engine_libreoffice_forced_without_lo_raises():
    """An operator who explicitly opts into LibreOffice must learn about a
    misconfigured worker — silent fallback would mask the wrong image
    being deployed."""
    with pytest.raises(RuntimeError, match="is on PATH"):
        _resolve_office_engine("libreoffice", docx_is_complex=False, soffice_available=False)


def test_engine_libreoffice_forced_with_lo_picks_libreoffice():
    engine, warnings = _resolve_office_engine(
        "libreoffice", docx_is_complex=False, soffice_available=True
    )
    assert engine == "libreoffice"
    assert warnings == []


def test_engine_mammoth_forced_ignores_complexity():
    engine, warnings = _resolve_office_engine(
        "mammoth", docx_is_complex=True, soffice_available=True
    )
    assert engine == "mammoth"
    assert warnings == []


def test_engine_value_is_case_insensitive():
    engine, _ = _resolve_office_engine("Mammoth", docx_is_complex=True, soffice_available=True)
    assert engine == "mammoth"
    engine, _ = _resolve_office_engine("LIBREOFFICE", docx_is_complex=False, soffice_available=True)
    assert engine == "libreoffice"


# ── DocxToPdfConverter routing (no actual rendering) ─────────────────────────


def test_converter_routes_complex_doc_to_libreoffice(tmp_path, monkeypatch):
    """When the detector flags complexity AND soffice is reachable,
    DocxToPdfConverter must invoke ``_convert_via_libreoffice`` and skip
    mammoth — verified by stubbing both functions."""
    from app.converters import document as doc_mod

    simple = _make_simple_docx(tmp_path / "base.docx")
    complex_docx = _inject_part(simple, tmp_path / "complex.docx", "word/footnotes.xml", b"<x/>")

    libreoffice_called = []
    mammoth_called = []

    def fake_lo(input_path, output_path, timeout_s):
        libreoffice_called.append((input_path, output_path, timeout_s))
        output_path.write_bytes(b"%PDF-1.4 fake")

    def fake_mammoth(input_path, output_path):
        mammoth_called.append((input_path, output_path))
        output_path.write_bytes(b"%PDF-1.4 fake")
        return False

    monkeypatch.setattr(doc_mod, "_convert_via_libreoffice", fake_lo)
    monkeypatch.setattr(doc_mod, "_convert_via_mammoth", fake_mammoth)
    monkeypatch.setattr(doc_mod, "_soffice_available", lambda: True)
    monkeypatch.setattr(doc_mod.shutil, "which", lambda name: "/usr/bin/soffice")

    converter = doc_mod.DocxToPdfConverter()
    out = tmp_path / "out.pdf"
    converter.convert(complex_docx, out)

    assert libreoffice_called, "LibreOffice path should have been invoked"
    assert not mammoth_called, "mammoth must not run when LibreOffice succeeds"
    assert converter.engine_used == "libreoffice"
    assert converter.warnings == []


def test_converter_falls_back_to_mammoth_with_warnings_when_no_soffice(tmp_path, monkeypatch):
    """Complex DOCX + no soffice on PATH → mammoth fallback + warnings."""
    from app.converters import document as doc_mod

    simple = _make_simple_docx(tmp_path / "base.docx")
    complex_docx = _inject_part(simple, tmp_path / "complex.docx", "word/header1.xml", b"<x/>")

    def fake_mammoth(input_path, output_path):
        output_path.write_bytes(b"%PDF-1.4 fake")
        return False  # had_warnings — not the same as the routing warnings

    monkeypatch.setattr(doc_mod, "_convert_via_mammoth", fake_mammoth)
    monkeypatch.setattr(doc_mod, "_soffice_available", lambda: False)

    converter = doc_mod.DocxToPdfConverter()
    out = tmp_path / "out.pdf"
    converter.convert(complex_docx, out)

    assert converter.engine_used == "mammoth"
    assert "engine=mammoth_fallback" in converter.warnings
    assert "reason=soffice_unavailable" in converter.warnings
    # The forced-mammoth path also names the simplified features so the
    # client can show a precise notice rather than a generic
    # "fidelity=reduced".
    assert any(w.startswith("simplified=headers") for w in converter.warnings)


def test_converter_runtime_error_in_auto_falls_back(tmp_path, monkeypatch):
    """If LibreOffice is on PATH but soffice exits non-zero (sick worker,
    broken profile, OOM), ``auto`` mode should still deliver a PDF via the
    mammoth fallback — at the cost of a routing warning."""
    from app.converters import document as doc_mod

    simple = _make_simple_docx(tmp_path / "base.docx")
    complex_docx = _inject_part(simple, tmp_path / "complex.docx", "word/footnotes.xml", b"<x/>")

    def fake_lo(input_path, output_path, timeout_s):
        raise RuntimeError("soffice exited with status 1")

    def fake_mammoth(input_path, output_path):
        output_path.write_bytes(b"%PDF-1.4 fake")
        return False

    monkeypatch.setattr(doc_mod, "_convert_via_libreoffice", fake_lo)
    monkeypatch.setattr(doc_mod, "_convert_via_mammoth", fake_mammoth)
    monkeypatch.setattr(doc_mod, "_soffice_available", lambda: True)

    converter = doc_mod.DocxToPdfConverter()
    out = tmp_path / "out.pdf"
    converter.convert(complex_docx, out)

    assert converter.engine_used == "mammoth"
    assert "engine=mammoth_fallback" in converter.warnings
    assert "reason=soffice_runtime_error" in converter.warnings


def test_converter_runtime_error_propagates_under_forced_libreoffice(tmp_path, monkeypatch):
    """When the operator pins ``office_engine=libreoffice``, a runtime
    failure must propagate — silent fallback would hide a misconfigured
    image. The mammoth fallback is reserved for ``auto`` mode."""
    from app.converters import document as doc_mod
    from app.core.config import settings

    simple = _make_simple_docx(tmp_path / "simple.docx")

    def fake_lo(input_path, output_path, timeout_s):
        raise RuntimeError("soffice exited with status 1")

    monkeypatch.setattr(doc_mod, "_convert_via_libreoffice", fake_lo)
    monkeypatch.setattr(doc_mod, "_soffice_available", lambda: True)
    monkeypatch.setattr(settings, "office_engine", "libreoffice")

    converter = doc_mod.DocxToPdfConverter()
    out = tmp_path / "out.pdf"
    with pytest.raises(RuntimeError):
        converter.convert(simple, out)
