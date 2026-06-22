# SPDX-License-Identifier: AGPL-3.0-or-later
"""Tests for the AI operations HTTP controller (app/api/routes/ai.py).

Covers the gates (inert-without-config 503, paid-tier 403), the two-phase
detect/apply flow, fail-closed behaviour, and — importantly — that responses
stay credit-denominated (no model id / token count / euro cost leaks, so the
margin stays opaque).

The happy-path tests enable the feature and (test-only) add ``anonymous`` to
the eligible tiers so they exercise the route without spinning up the full
auth/DB stack. Production keeps ``anonymous`` out (paid-only).
"""

import pytest

from app.core.config import settings

# A text note carrying structured PII the deterministic engine catches.
PII_NOTE = b"Kontakt: max.mustermann@beispiel.de, IBAN DE89 3704 0044 0532 0130 00."

# Tokens that would reveal cost structure / margin if they ever leaked.
_FORBIDDEN_LEAK_TOKENS = ("model", "token", "haiku", "sonnet", "claude", "eur", "cost_", "presidio")


def _set_ai(**overrides):
    """Patch the settings singleton's __dict__ (same pattern as conftest)."""
    d = settings.__dict__
    saved = {k: d.get(k) for k in overrides}
    d.update(overrides)
    return saved


def _restore(saved):
    settings.__dict__.update(saved)


@pytest.fixture
def ai_enabled():
    """Feature on + anonymous allowed (test-only) so the route runs without auth."""
    saved = _set_ai(
        ai_operations_enabled=True,
        ai_eligible_tiers="anonymous,pro,business,enterprise",
    )
    yield
    _restore(saved)


# ---------------------------------------------------------------------------
# gates
# ---------------------------------------------------------------------------


def test_detect_503_when_disabled(client, auth_headers):
    # Default: ai_operations_enabled is False → inert.
    resp = client.post(
        "/api/v1/ai/redact/detect",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
    )
    assert resp.status_code == 503
    assert resp.headers.get("X-FileMorph-Error-Code") == "ai_unavailable"


def test_apply_503_when_disabled(client, auth_headers):
    resp = client.post(
        "/api/v1/ai/redact/apply",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
    )
    assert resp.status_code == 503


def test_detect_is_free_for_ineligible_tier(client, auth_headers):
    # detect is the free findings preview — open even when the caller's tier is
    # NOT eligible (the hook that leads to the paid apply).
    saved = _set_ai(ai_operations_enabled=True, ai_eligible_tiers="pro,business,enterprise")
    try:
        resp = client.post(
            "/api/v1/ai/redact/detect",
            headers=auth_headers,
            files={"file": ("note.txt", PII_NOTE, "text/plain")},
        )
    finally:
        _restore(saved)
    assert resp.status_code == 200
    assert resp.json()["count"] >= 1


def test_apply_requires_paid_tier(client, auth_headers):
    # apply stays paid-gated: an ineligible/anonymous caller → 403.
    saved = _set_ai(ai_operations_enabled=True, ai_eligible_tiers="pro,business,enterprise")
    try:
        resp = client.post(
            "/api/v1/ai/redact/apply",
            headers=auth_headers,
            files={"file": ("note.txt", PII_NOTE, "text/plain")},
        )
    finally:
        _restore(saved)
    assert resp.status_code == 403
    assert resp.headers.get("X-FileMorph-Error-Code") == "ai_plan_required"


# ---------------------------------------------------------------------------
# detect (phase 1)
# ---------------------------------------------------------------------------


def test_detect_finds_pii(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/detect",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
    )
    assert resp.status_code == 200
    body = resp.json()
    types = {f["entity_type"] for f in body["findings"]}
    assert "EMAIL" in types
    assert "IBAN" in types
    assert body["count"] == len(body["findings"])
    assert body["credits_estimate"] == settings.ai_credit_cost_redact


def test_detect_response_has_no_cost_leak(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/detect",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
    )
    blob = resp.text.lower()
    for tok in _FORBIDDEN_LEAK_TOKENS:
        assert tok not in blob, f"cost-structure leak: {tok!r} appeared in detect response"


def test_detect_scope_filter(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/detect",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
        data={"entity_types": "EMAIL"},
    )
    assert resp.status_code == 200
    types = {f["entity_type"] for f in resp.json()["findings"]}
    assert types == {"EMAIL"}  # IBAN out of scope


def test_detect_unknown_entity_type_400(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/detect",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
        data={"entity_types": "BANANA"},
    )
    assert resp.status_code == 400
    assert resp.headers.get("X-FileMorph-Error-Code") == "unknown_entity_type"


def test_detect_binary_415(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/detect",
        headers=auth_headers,
        files={"file": ("blob.bin", b"\xff\xfe\xfa\x01\x02not utf8", "application/octet-stream")},
    )
    assert resp.status_code == 415
    assert resp.headers.get("X-FileMorph-Error-Code") == "unsupported_format"


# ---------------------------------------------------------------------------
# apply (phase 2)
# ---------------------------------------------------------------------------


def test_apply_redacts_and_verifies(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/apply",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
    )
    assert resp.status_code == 200
    out = resp.content.decode("utf-8")
    assert "max.mustermann@beispiel.de" not in out
    assert "DE89" not in out
    assert "[EMAIL]" in out and "[IBAN]" in out
    assert int(resp.headers["X-FileMorph-AI-Entities-Redacted"]) >= 2
    assert resp.headers["X-FileMorph-AI-Credits-Cost"] == str(settings.ai_credit_cost_redact)
    assert "note.redacted.txt" in resp.headers.get("Content-Disposition", "")


def test_apply_mask_mode(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/apply",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
        data={"mode": "mask"},
    )
    assert resp.status_code == 200
    out = resp.content.decode("utf-8")
    assert "max.mustermann@beispiel.de" not in out
    assert "*" in out


def test_apply_500_on_verification_failure(client, auth_headers, ai_enabled, monkeypatch):
    """Fail-closed: if redaction can't be verified, return 500 and NO file."""
    import app.ee.ai_ops as ai_ops
    from app.ee.ai_ops import redaction as rd

    def _fake_redact(text, types=None, mode="replace"):
        return rd.RedactionResult(
            text=text, spans=[], entities_redacted=0, verification_passed=False, residual=[]
        )

    monkeypatch.setattr(ai_ops, "redact_text", _fake_redact)
    resp = client.post(
        "/api/v1/ai/redact/apply",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
    )
    assert resp.status_code == 500
    assert resp.headers.get("X-FileMorph-Error-Code") == "redaction_verification_failed"
    assert "max.mustermann@beispiel.de" not in resp.text  # no half-redacted content leaks


def test_apply_unknown_mode_400(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/apply",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
        data={"mode": "blackout"},
    )
    assert resp.status_code == 400


def test_apply_response_has_no_cost_leak(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/apply",
        headers=auth_headers,
        files={"file": ("note.txt", PII_NOTE, "text/plain")},
    )
    header_blob = " ".join(f"{k}:{v}" for k, v in resp.headers.items()).lower()
    for tok in _FORBIDDEN_LEAK_TOKENS:
        assert tok not in header_blob, f"cost-structure leak in headers: {tok!r}"


# ---------------------------------------------------------------------------
# binary formats (DOCX / XLSX) + unsupported (PDF) + corrupt
# ---------------------------------------------------------------------------

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _make_docx(lines: list[str]) -> bytes:
    import io

    from docx import Document

    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_text(data: bytes) -> str:
    import io

    from docx import Document

    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def _make_xlsx(values: list[str]) -> bytes:
    import io

    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    for i, v in enumerate(values, start=1):
        ws.cell(row=i, column=1, value=v)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _xlsx_values(data: bytes) -> list[str]:
    import io

    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data))
    return [
        cell.value
        for ws in wb.worksheets
        for row in ws.iter_rows()
        for cell in row
        if isinstance(cell.value, str)
    ]


def test_detect_docx(client, auth_headers, ai_enabled):
    data = _make_docx(["Kontakt: max.mustermann@beispiel.de", "IBAN DE89 3704 0044 0532 0130 00"])
    resp = client.post(
        "/api/v1/ai/redact/detect",
        headers=auth_headers,
        files={"file": ("brief.docx", data, _DOCX_MIME)},
    )
    assert resp.status_code == 200
    types = {f["entity_type"] for f in resp.json()["findings"]}
    assert {"EMAIL", "IBAN"} <= types


def test_apply_docx_redacts(client, auth_headers, ai_enabled):
    data = _make_docx(["Kontakt: max.mustermann@beispiel.de", "Konto DE89 3704 0044 0532 0130 00"])
    resp = client.post(
        "/api/v1/ai/redact/apply",
        headers=auth_headers,
        files={"file": ("brief.docx", data, _DOCX_MIME)},
    )
    assert resp.status_code == 200
    assert _DOCX_MIME in resp.headers["content-type"]
    out_text = _docx_text(resp.content)
    assert "max.mustermann@beispiel.de" not in out_text
    assert "DE89" not in out_text
    assert "[EMAIL]" in out_text and "[IBAN]" in out_text
    assert int(resp.headers["X-FileMorph-AI-Entities-Redacted"]) >= 2
    assert "brief.redacted.docx" in resp.headers.get("Content-Disposition", "")


def test_detect_xlsx(client, auth_headers, ai_enabled):
    data = _make_xlsx(["max.mustermann@beispiel.de", "DE89370400440532013000", "kein PII"])
    resp = client.post(
        "/api/v1/ai/redact/detect",
        headers=auth_headers,
        files={"file": ("tab.xlsx", data, _XLSX_MIME)},
    )
    assert resp.status_code == 200
    types = {f["entity_type"] for f in resp.json()["findings"]}
    assert {"EMAIL", "IBAN"} <= types


def test_apply_xlsx_redacts(client, auth_headers, ai_enabled):
    data = _make_xlsx(["max.mustermann@beispiel.de", "DE89370400440532013000", "kein PII"])
    resp = client.post(
        "/api/v1/ai/redact/apply",
        headers=auth_headers,
        files={"file": ("tab.xlsx", data, _XLSX_MIME)},
    )
    assert resp.status_code == 200
    joined = " ".join(_xlsx_values(resp.content))
    assert "max.mustermann@beispiel.de" not in joined
    assert "DE89370400440532013000" not in joined
    assert "kein PII" in joined  # clean cell untouched


def test_pdf_returns_415(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/detect",
        headers=auth_headers,
        files={"file": ("scan.pdf", b"%PDF-1.4\n1 0 obj\n", "application/pdf")},
    )
    assert resp.status_code == 415
    assert resp.headers.get("X-FileMorph-Error-Code") == "unsupported_format"


def test_corrupt_docx_400(client, auth_headers, ai_enabled):
    resp = client.post(
        "/api/v1/ai/redact/apply",
        headers=auth_headers,
        files={"file": ("broken.docx", b"this is not a real docx file at all", _DOCX_MIME)},
    )
    assert resp.status_code == 400
    assert resp.headers.get("X-FileMorph-Error-Code") == "document_unreadable"
