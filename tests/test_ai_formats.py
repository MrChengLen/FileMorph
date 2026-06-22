# SPDX-License-Identifier: AGPL-3.0-or-later
"""Format-layer redaction regression guard (app/ee/ai_ops/formats.py).

The headline test class is the C1 fix: PII hidden in DOCX/XLSX metadata, comments
and sheet names used to survive redaction while the verifier falsely reported
success (it re-scanned only the nodes it had touched). These tests redact such
documents and then **unzip the produced package and grep every part for the live
value** — the same extraction the audit used to break it — so a regression that
re-introduces the leak fails loudly. They also pin the honesty of the fail-closed
gate and that the rewrite did not start refusing ordinary numeric spreadsheets.
"""

from __future__ import annotations

import io
import zipfile

import pytest

from app.ee.ai_ops.formats import (
    DocumentReadError,
    _mop_up_package,
    _package_residual,
    detect_docx,
    detect_xlsx,
    redact_docx,
    redact_xlsx,
)

IBAN = "DE89370400440532013000"  # canonical valid German IBAN
EMAIL = "leak@example.com"
CARD = "4111111111111111"  # Luhn-valid Visa test number
SECRETS = (IBAN, EMAIL, CARD)


def _parts_contain(data: bytes, needle: str) -> bool:
    """True if `needle` appears verbatim in ANY part of the OOXML package."""
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return any(needle in zf.read(n).decode("utf-8", "replace") for n in zf.namelist())


def _make_docx(*, author="", title="", comments="", paragraphs=(), table_cells=()):
    from docx import Document

    doc = Document()
    doc.core_properties.author = author
    doc.core_properties.title = title
    doc.core_properties.comments = comments
    for line in paragraphs:
        doc.add_paragraph(line)
    if table_cells:
        t = doc.add_table(rows=len(table_cells), cols=1)
        for i, val in enumerate(table_cells):
            t.rows[i].cells[0].text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx(*, sheet_title=None, creator="", title="", cells=()):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    if sheet_title is not None:
        ws.title = sheet_title
    wb.properties.creator = creator
    wb.properties.title = title
    for i, val in enumerate(cells, start=1):
        ws.cell(row=i, column=1, value=val)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── C1: PII in metadata / comments / sheet names must NOT survive ───────────


def test_docx_metadata_pii_does_not_survive():
    """The auditor's primary finding: IBAN/email in core properties survived."""
    data = _make_docx(author=IBAN, title=EMAIL, comments=CARD, paragraphs=["clean body"])
    res = redact_docx(data, None, "replace")
    assert res.verification_passed
    assert res.data
    for secret in SECRETS:
        assert not _parts_contain(res.data, secret), f"{secret} survived in a package part"


def test_docx_body_and_table_pii_redacted():
    data = _make_docx(
        paragraphs=[f"Mail {EMAIL}", "clean"],
        table_cells=[f"Konto {IBAN}"],
    )
    res = redact_docx(data, None, "replace")
    assert res.verification_passed
    assert not _parts_contain(res.data, EMAIL)
    assert not _parts_contain(res.data, IBAN)


def test_redacted_docx_reopens_as_valid_document():
    from docx import Document

    data = _make_docx(author=IBAN, paragraphs=[f"Mail {EMAIL}"])
    res = redact_docx(data, None, "replace")
    Document(io.BytesIO(res.data))  # must not raise


def test_xlsx_sheet_name_and_property_pii_does_not_survive():
    data = _make_xlsx(sheet_title="Kunde leak@x.de", creator=EMAIL, title=IBAN, cells=["clean"])
    res = redact_xlsx(data, None, "replace")
    assert res.verification_passed
    assert res.data
    for secret in (EMAIL, IBAN, "leak@x.de"):
        assert not _parts_contain(res.data, secret), f"{secret} survived in a package part"


def test_xlsx_cell_pii_redacted_clean_cell_untouched():
    data = _make_xlsx(cells=[f"pay {IBAN}", "kein PII", f"card {CARD}"])
    res = redact_xlsx(data, None, "replace")
    assert res.verification_passed
    assert not _parts_contain(res.data, IBAN)
    assert not _parts_contain(res.data, CARD)
    assert _parts_contain(res.data, "kein PII")  # clean content preserved


def test_redacted_xlsx_reopens_as_valid_workbook():
    import openpyxl

    data = _make_xlsx(creator=EMAIL, cells=[f"pay {IBAN}"])
    res = redact_xlsx(data, None, "replace")
    openpyxl.load_workbook(io.BytesIO(res.data))  # must not raise


# ── Honesty of the fail-closed gate ─────────────────────────────────────────


def test_gate_flags_pii_the_mop_up_skips():
    """workbook.xml is skipped by the mop-up (sheet-name char constraints); the
    gate MUST still detect PII there so the caller fails closed — this is the
    exact false-pass the C1 fix removes."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", "<types/>")
        zf.writestr(
            "xl/workbook.xml", f'<workbook><sheets><sheet name="{IBAN}"/></sheets></workbook>'
        )
    mopped, _extra = _mop_up_package(buf.getvalue(), None, "replace")
    assert _package_residual(mopped, None) > 0


def test_mop_up_cleans_an_unreachable_text_part():
    """A comments part the high-level API cannot reach is cleaned by the mop-up."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", "<types/>")
        zf.writestr(
            "word/comments.xml", f"<w:comments><w:comment>{IBAN} {EMAIL}</w:comment></w:comments>"
        )
    mopped, extra = _mop_up_package(buf.getvalue(), None, "replace")
    assert extra >= 2
    assert _package_residual(mopped, None) == 0


def test_embedded_object_is_refused():
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", "<types/>")
        zf.writestr("word/document.xml", "<w:document/>")
        zf.writestr("word/embeddings/oleObject1.bin", b"\x00\x01binary")
    with pytest.raises(DocumentReadError):
        redact_docx(buf.getvalue(), None, "replace")


# ── C1 (re-opened by audit): PII in XML ATTRIBUTE values must not survive ────


def _inject_part(data: bytes, part_name: str, new_bytes: bytes) -> bytes:
    """Rewrite one part of an OOXML package with new bytes."""
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(data)) as zin, zipfile.ZipFile(out, "w") as zout:
        for info in zin.infolist():
            content = new_bytes if info.filename == part_name else zin.read(info.filename)
            zout.writestr(info, content)
    return out.getvalue()


def test_gate_detects_pii_in_an_attribute_value():
    """The audit's re-opened finding: PII hidden in an attribute of a part that is
    not .rels/workbook.xml. The gate must see it (fail-closed)."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", "<types/>")
        zf.writestr(
            "word/settings.xml", f'<w:settings><w:docVar w:val="{EMAIL} {IBAN}"/></w:settings>'
        )
    assert _package_residual(buf.getvalue(), None) > 0


def test_gate_no_false_positive_on_rsid_attributes():
    """rsid / paraId structural IDs are 8 hex chars and trip the checksum-less
    phone detector; the attribute scan must skip that shape so a clean doc isn't
    refused (this is the false-positive the markup-scan approach caused)."""
    rsids = "".join(
        f'<w:p w:rsidR="0{i:07d}"><w:r><w:t>row {i}</w:t></w:r></w:p>' for i in range(40)
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", "<types/>")
        zf.writestr("word/document.xml", f"<w:document><w:body>{rsids}</w:body></w:document>")
    assert _package_residual(buf.getvalue(), None) == 0


def test_docx_settings_docvar_attribute_pii_does_not_survive():
    data = _make_docx(paragraphs=["clean"])
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        settings = zf.read("word/settings.xml").decode("utf-8")
    injected = settings.replace(
        "</w:settings>",
        f'<w:docVars><w:docVar w:name="x" w:val="{EMAIL} {IBAN}"/></w:docVars></w:settings>',
    )
    data = _inject_part(data, "word/settings.xml", injected.encode("utf-8"))
    res = redact_docx(data, None, "replace")
    if res.verification_passed:
        assert not _parts_contain(res.data, EMAIL)
        assert not _parts_contain(res.data, IBAN)
    else:
        assert res.data == b""  # refused, fail-closed — never a false pass


def test_xlsx_custom_property_pii_does_not_survive():
    import openpyxl
    from openpyxl.packaging.custom import StringProperty

    wb = openpyxl.Workbook()
    wb.active["A1"] = "clean"
    wb.custom_doc_props.append(StringProperty(name=EMAIL, value=IBAN))  # PII in name + value
    buf = io.BytesIO()
    wb.save(buf)
    res = redact_xlsx(buf.getvalue(), None, "replace")
    assert res.verification_passed
    assert not _parts_contain(res.data, EMAIL)
    assert not _parts_contain(res.data, IBAN)


# ── No regression: numeric spreadsheets must not be falsely refused ─────────


def test_numeric_xlsx_is_not_falsely_refused():
    """Scanning markup (not text) would match RSIDs / cell numbers as phones and
    cards and refuse clean files. Text-only scanning must let a numeric sheet
    through untouched."""
    data = _make_xlsx(cells=["Menge", 12, 19.99, 7, 2024, "Summe", 1000])
    res = redact_xlsx(data, None, "replace")
    assert res.verification_passed
    assert res.data
    assert res.entities_redacted == 0


# ── /detect preview surfaces metadata findings (honest preview) ─────────────


def test_detect_docx_surfaces_metadata_pii():
    data = _make_docx(author=IBAN, paragraphs=["clean"])
    findings = detect_docx(data, None)
    assert any(f["value"] == IBAN for f in findings)


def test_detect_xlsx_surfaces_sheet_name_pii():
    data = _make_xlsx(sheet_title="Kunde leak@x.de", cells=["clean"])
    findings = detect_xlsx(data, None)
    assert any(f["entity_type"] == "EMAIL" for f in findings)
